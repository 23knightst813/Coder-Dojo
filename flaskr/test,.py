# Creating a combined list of variables, ensuring no duplicates by choosing preferred names and descriptions.
from docx import Document


# Define all variables from both lists, removing duplicates
combined_variables_unique = [
    ["Variable Name", "Data Type", "Description", "Location"],
    ["app", "Flask object", "The Flask application instance.", "app.py"],
    ["app.secret_key", "str", "Secret key for session management.", "app.py"],
    ["email", "str", "User's email address.", "app.py, auth.py"],
    ["password", "str", "User's chosen password.", "app.py, auth.py, Register System"],
    ["first_name", "str", "User's first name.", "app.py"],
    ["last_name", "str", "User's last name.", "app.py"],
    ["password_confirm", "str", "Confirmation of user's password.", "app.py"],
    ["EMAIL", "str", "User's email address (used in login).", "app.py"],  # Keeping this for reference.
    ["response", "Response object", "HTTP response to be sent to the client.", "app.py"],
    ["conn", "sqlite3.Connection", "Connection object to the SQLite database.", "db.py, auth.py"],
    ["cur", "sqlite3.Cursor", "Cursor object for executing SQL commands.", "db.py, auth.py"],
    ["queries", "list of str", "List of SQL queries to create tables.", "db.py"],
    ["query", "str", "SQL query string for inserting/selecting data.", "db.py, auth.py"],
    ["Creation_Date", "datetime.datetime", "Timestamp of user account creation.", "db.py"],
    ["user", "tuple or None", "Fetched user data from the database.", "auth.py"],
    ["session", "dict", "Flask session object for storing user session data.", "Used across multiple routes"],
    ["request", "Request object", "Flask request object containing HTTP request data.", "app.py"],
    ["flash", "Function", "Flask function to flash messages to the next request.", "app.py, auth.py"],
    ["render_template", "Function", "Flask function to render HTML templates.", "app.py"],
    ["redirect", "Function", "Flask function to redirect to a different route.", "app.py"],
    ["make_response", "Function", "Flask function to create a response object.", "app.py"],
    ["url_for", "Function", "Flask function to build a URL to a specific function.", "app.py"],
    ["user_email", "String", "Email address of the user", "Admin Dashboard"],
    ["user_time", "DateTime", "Time associated with the user's action", "Admin Dashboard"],
    ["bookings", "Array", "List of bookings", "Admin Dashboard"],
    ["new_record", "Boolean", "Flag for new record addition", "Admin Dashboard"],
    ["current_bookings", "Array", "List of current bookings", "Admin Dashboard"],
    ["booking_id", "Integer", "Unique identifier for each booking", "Admin Dashboard, Booking System"],
    ["b_day", "Date", "Day of the booking", "Booking System"],
    ["b_time", "Time", "Time of the booking", "Booking System"],
    ["b_name", "String", "Name associated with the booking", "Booking System"],
    ["b_location", "String", "Location of the booking", "Booking System"],
    ["b_activity", "String", "Activity for the booking", "Booking System"],
    ["overflow_count", "Integer", "Number of overflow bookings", "Booking System"],
    ["user_password", "String", "Password for user login", "Sub Routines"],
    ["hash_user_password", "String", "Hashed version of the user's password", "Sub Routines"],
    ["user_id", "Integer", "Unique user identifier", "Sub Routines"],
    ["user_login", "Boolean", "User login status", "Sub Routines"],
    ["session_permanent", "Boolean", "Session persistence flag", "Sub Routines"],
    ["user_name", "String", "User's full name", "Register System"],
    ["booking_count", "Integer", "Total number of bookings", "General"]
]

# Create a new Document for the entire combined and deduplicated variables list
doc = Document()

# Add main heading
doc.add_heading("Comprehensive Variables List", level=1)

# Add table
table = doc.add_table(rows=1, cols=len(combined_variables_unique[0]))
hdr_cells = table.rows[0].cells
for i, header in enumerate(combined_variables_unique[0]):
    hdr_cells[i].text = header

# Populate the table with the unique variables
for row_data in combined_variables_unique[1:]:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Save the document
file_path_comprehensive_variables = "flaskr/templates/Comprehensive_Variables_List.docx"
doc.save(file_path_comprehensive_variables)
file_path_comprehensive_variables
